import os
from pathlib import Path
from docx import Document

from scripts.embed_final_rfp import extract_qa_from_docx


def create_sample_docx(path: Path):
    doc = Document()
    doc.add_paragraph("What is your name?")
    doc.add_paragraph("My name is Bot.")
    doc.add_paragraph("What do you do?")
    doc.add_paragraph("I automate tasks.")
    doc.save(path)


def test_extract_qa_from_docx(tmp_path):
    file_path = tmp_path / "sample.docx"
    create_sample_docx(file_path)
    qa_pairs = extract_qa_from_docx(file_path)
    assert qa_pairs == [
        {"question": "What is your name?", "answer": "My name is Bot."},
        {"question": "What do you do?", "answer": "I automate tasks."},
    ]

