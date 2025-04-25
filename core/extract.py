# import os
# from docx import Document


# def extract_questions_from_docx(file_path):
#     doc = Document(file_path)
#     questions = []
#     for para in doc.paragraphs:
#         text = para.text.strip()
#         if text.endswith("?") and len(text.split()) > 3:
#             questions.append(text)
#         return questions

from docx import Document


def extract_questions_from_docx(file_path):
    doc = Document(file_path)
    questions = []

    print("\n All paragraphs found in docx:")
    for para in doc.paragraphs:
        text = para.text.strip()
        print(f"• {text}")
        if text.endswith("?") and len(text.split()) > 3:
            questions.append(text)

    print(f"\nExtracted {len(questions)} question(s)")
    return questions
