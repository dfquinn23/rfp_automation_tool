import json
import os
from docx import Document
from qdrant_client import QdrantClient
from qdrant_client.http.models import SearchParams
from generate_drafts import generate_draft_answer
from qa_generation import generate_reviewable_draft
from dotenv import load_dotenv
import requests

load_dotenv()

# Config
QDRANT_API_KEY = os.getenv("QDRANT_API_KEY")
QDRANT_CLUSTER_URL = "https://295f06f4-25f8-4197-bd5e-2227d256a1f1.us-east4-0.gcp.cloud.qdrant.io:6333"
COLLECTION_NAME = "past_rfp_answers"
EXTRACTED_QUESTIONS_PATH = "output/extracted_questions.json"
OUTPUT_DOCX_PATH = "output/generated_rfp_draft.docx"
OLLAMA_MODEL = "llama3"

# Prompt C instruction
PROMPT_C_INSTRUCTION = (
    "Use the first answer as a foundation."
    "Improve or expand it using insight from the others and create a concise, professional response"
)

# Connect to qdrant
client = QdrantClient(
    url=QDRANT_CLUSTER_URL,
    api_key=QDRANT_API_KEY
)

# load new RFP questions
with open(EXTRACTED_QUESTIONS_PATH, "r", encoding="utf-8") as f:
    all_questions = json.load(f)

# Create new word doc
doc = Document()
doc.add_heading("RFP Draft Responses", level=1)

# Loop through questions & genertate answers
for filename, questions in all_questions.items():
    doc.add_heading(f"File: {filename}", level=2)

    for i, question in enumerate(questions, 1):
        print(f"🔍 Q{i}: {question}")

        # embed question
        embedding = requests.post(
            "http://localhost:11434/api/embeddings",
            json={"model": "nomic-embed-text", "prompt": question}
        ).json()["embedding"]

        # search qdrant
        results = client.search(
            collection_name=COLLECTION_NAME,
            query_vector=embedding,
            limit=3,
            with_payload=True,
            search_params=SearchParams(hnsw_ef=128)
        )
        top_answers = [r.payload["answer"] for r in results]

        # generate draft
        draft = generate_draft_answer(
            question, top_answers, model=OLLAMA_MODEL, extra_instruction=PROMPT_C_INSTRUCTION)

        # add to document
        result = generate_reviewable_draft(question)

        doc.add_paragraph(f"Q{i}. {result['question']}", style="List Number")
        doc.add_paragraph(result['draft'])
        doc.add_paragraph("")  # spacing

doc.save(OUTPUT_DOCX_PATH)
print(f"\n✅ Draft document saved to: {OUTPUT_DOCX_PATH}")
