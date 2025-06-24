# run_pipeline.py
import os
import argparse
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from core.config import OUTPUT_DIR, REVIEW_SCORE_THRESHOLD
from core.extract import extract_questions_from_docx
# from core.embed import embed_text
from core.generate import get_embedding
from core.search import search_qdrant
from core.generate import generate_draft_answer
from core.logger import log_result
from pathlib import Path
from dotenv import load_dotenv
# This automatically finds the project root and loads the .env file.
project_root = Path(__file__).parent.parent
dotenv_path = project_root / ".env"
load_dotenv(dotenv_path=dotenv_path)


def run_pipeline(input_path):
    print(f"\n[-->] Loading RFP: {input_path}")
    questions = extract_questions_from_docx(input_path)
    if not questions:
        print("X No valid questions found.")
        return

    print(
        f"Extracted {len(questions)} questions. Starting draft generation...\n")

    os.makedirs(OUTPUT_DIR, exist_ok=True)
    full_doc = Document()
    full_doc.add_heading("RFP Draft Responses", level=1)

    review_doc = Document()
    review_doc.add_heading("[!] Needs Review", level=1)

   # ADD THIS NEW, CORRECTED BLOCK in run_pipeline.py
for i, question in enumerate(questions, 1):
    print(f"Q{i}: {question}")

    vector = get_embedding(question)
    results = search_qdrant(vector)  # results is a list of ScoredPoint objects

    # Let the generate_draft_answer function handle the results directly
    draft = generate_draft_answer(question, results)

    # Check the score of the top result to see if it needs review
    top_score = results[0].score if results else 0.0
    needs_review = top_score < REVIEW_SCORE_THRESHOLD

    # Prepend the "Needs review" flag here if necessary
    if not results or needs_review:
        draft = f"[⚠ Needs review | Top Score: {top_score:.2f}]\n{draft}"

    # Log the results for auditing
    log_result({
        "question": question,
        "top_score": top_score,
        "needs_review": needs_review,
        # Log the full payload
        "top_answers_payload": [r.payload for r in results],
        "draft": draft
    })

    # Add the question and the generated draft to the document
    p_q = full_doc.add_paragraph()
    run_q = p_q.add_run(f"Q{i}: {question}")
    run_q.bold = True
    run_q.font.size = Pt(11)
    p_q.space_after = Pt(6)

    p_a = full_doc.add_paragraph(draft)
    p_a.space_after = Pt(14)

    # Add to the review document if needed
    if needs_review:
        p_q_r = review_doc.add_paragraph()
        run_q_r = p_q_r.add_run(f"Q{i}: {question}")
        run_q_r.bold = True
        run_q_r.font.size = Pt(11)
        p_q_r.space_after = Pt(6)

        p_a_r = review_doc.add_paragraph(draft)
        p_a_r.space_after = Pt(14)

        p_q = full_doc.add_paragraph()
        run_q = p_q.add_run(f"Q{i}: {question}")
        run_q.bold = True
        run_q.font.size = Pt(11)
        p_q.space_after = Pt(6)

        p_a = full_doc.add_paragraph(draft)
        p_a.space_after = Pt(14)

        if needs_review:
            p_q_r = review_doc.add_paragraph()
            run_q_r = p_q_r.add_run(f"Q{i}: {question}")
            run_q_r.bold = True
            run_q_r.font.size = Pt(11)
            p_q_r.space_after = Pt(6)

            p_a_r = review_doc.add_paragraph(draft)
            p_a_r.space_after = Pt(14)

    full_path = os.path.join(OUTPUT_DIR, "generated_rfp_draft.docx")
    full_doc.save(full_path)
    print(f"Full draft saved to: {full_path}")

    if review_doc.paragraphs:
        review_path = os.path.join(OUTPUT_DIR, "low_confidence_rfp_draft.docx")
        review_doc.save(review_path)
        print(f"[!] Low-confidence draft saved to: {review_path}")


if __name__ == "__main__":
    parser = argparse.ArgumentParser(
        description="Run full RFP automation pipeline.")
    parser.add_argument("--rfp", required=True, help="Path to new RFP .docx")
    args = parser.parse_args()

    run_pipeline(args.rfp)
